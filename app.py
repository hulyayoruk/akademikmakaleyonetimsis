import streamlit as st
import requests
import io
import re
import math
from collections import Counter

import nltk

from pypdf import PdfReader
from deep_translator import GoogleTranslator
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

nltk.download("stopwords", quiet=True)
nltk.download("punkt", quiet=True)
nltk.download("punkt_tab", quiet=True)

from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize, word_tokenize

st.set_page_config(
    page_title="Akademik Makale Yönetim Sistemi",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="collapsed"
)

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Playfair+Display:wght@600;700&display=swap');

    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif;
    }

    .stApp {
        background: linear-gradient(135deg, #0f0c29, #302b63, #24243e);
        min-height: 100vh;
    }

    .main-header {
        text-align: center;
        padding: 2.5rem 0 1.5rem 0;
    }

    .main-header h1 {
        font-family: 'Playfair Display', serif;
        font-size: 2.6rem;
        font-weight: 700;
        color: #e8d5b7;
        letter-spacing: -0.5px;
        margin-bottom: 0.4rem;
        text-shadow: 0 2px 20px rgba(232,213,183,0.3);
    }

    .main-header p {
        color: #9b8ea0;
        font-size: 1rem;
        font-weight: 300;
        letter-spacing: 1px;
    }

    .divider-line {
        height: 1px;
        background: linear-gradient(90deg, transparent, #e8d5b7, transparent);
        margin: 0.5rem auto 2rem auto;
        width: 60%;
        opacity: 0.4;
    }

    .card {
        background: rgba(255,255,255,0.04);
        border: 1px solid rgba(232,213,183,0.15);
        border-radius: 14px;
        padding: 1.6rem;
        margin-bottom: 1.2rem;
        backdrop-filter: blur(10px);
        transition: border-color 0.3s ease;
    }

    .card:hover {
        border-color: rgba(232,213,183,0.35);
    }

    .card-title {
        font-family: 'Playfair Display', serif;
        font-size: 1.15rem;
        color: #e8d5b7;
        font-weight: 600;
        margin-bottom: 0.8rem;
        line-height: 1.4;
    }

    .card-meta {
        font-size: 0.78rem;
        color: #7c6f8a;
        margin-bottom: 0.6rem;
        letter-spacing: 0.3px;
    }

    .section-label {
        font-size: 0.72rem;
        font-weight: 600;
        letter-spacing: 1.5px;
        text-transform: uppercase;
        color: #c9a96e;
        margin-bottom: 0.35rem;
        margin-top: 0.9rem;
    }

    .text-block {
        font-size: 0.88rem;
        color: #b8aec4;
        line-height: 1.7;
        background: rgba(0,0,0,0.2);
        border-radius: 8px;
        padding: 0.8rem 1rem;
        border-left: 3px solid rgba(201,169,110,0.4);
    }

    .badge {
        display: inline-block;
        background: rgba(201,169,110,0.15);
        border: 1px solid rgba(201,169,110,0.3);
        color: #c9a96e;
        border-radius: 20px;
        padding: 0.2rem 0.7rem;
        font-size: 0.72rem;
        font-weight: 500;
        margin-right: 0.4rem;
        margin-bottom: 0.3rem;
    }

    .info-pill {
        display: inline-flex;
        align-items: center;
        gap: 0.3rem;
        background: rgba(255,255,255,0.05);
        border: 1px solid rgba(255,255,255,0.1);
        border-radius: 20px;
        padding: 0.25rem 0.8rem;
        font-size: 0.75rem;
        color: #9b8ea0;
        margin-right: 0.5rem;
        margin-bottom: 0.5rem;
    }

    .stButton > button {
        background: linear-gradient(135deg, #c9a96e, #a07840) !important;
        color: #1a1425 !important;
        border: none !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
        font-size: 0.82rem !important;
        padding: 0.45rem 1.1rem !important;
        transition: all 0.2s ease !important;
        letter-spacing: 0.3px !important;
    }

    .stButton > button:hover {
        transform: translateY(-1px) !important;
        box-shadow: 0 4px 15px rgba(201,169,110,0.4) !important;
    }

    .stDownloadButton > button {
        background: rgba(201,169,110,0.1) !important;
        color: #c9a96e !important;
        border: 1px solid rgba(201,169,110,0.4) !important;
        border-radius: 8px !important;
        font-weight: 500 !important;
        font-size: 0.8rem !important;
        padding: 0.4rem 1rem !important;
        transition: all 0.2s ease !important;
    }

    .stDownloadButton > button:hover {
        background: rgba(201,169,110,0.2) !important;
        border-color: rgba(201,169,110,0.7) !important;
    }

    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea {
        background: rgba(255,255,255,0.05) !important;
        border: 1px solid rgba(232,213,183,0.2) !important;
        border-radius: 8px !important;
        color: #e8d5b7 !important;
        font-family: 'Inter', sans-serif !important;
    }

    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus {
        border-color: rgba(201,169,110,0.6) !important;
        box-shadow: 0 0 0 2px rgba(201,169,110,0.15) !important;
    }

    .stSelectbox > div > div {
        background: rgba(255,255,255,0.05) !important;
        border: 1px solid rgba(232,213,183,0.2) !important;
        border-radius: 8px !important;
        color: #e8d5b7 !important;
    }

    .stSlider > div > div > div {
        color: #c9a96e !important;
    }

    .stFileUploader > div {
        background: rgba(255,255,255,0.03) !important;
        border: 2px dashed rgba(201,169,110,0.3) !important;
        border-radius: 12px !important;
    }

    .stTabs [data-baseweb="tab-list"] {
        background: rgba(255,255,255,0.03) !important;
        border-radius: 10px !important;
        padding: 0.3rem !important;
        gap: 0.3rem !important;
    }

    .stTabs [data-baseweb="tab"] {
        background: transparent !important;
        color: #9b8ea0 !important;
        border-radius: 8px !important;
        font-weight: 500 !important;
        font-size: 0.9rem !important;
        padding: 0.5rem 1.2rem !important;
        border: none !important;
    }

    .stTabs [aria-selected="true"] {
        background: rgba(201,169,110,0.2) !important;
        color: #e8d5b7 !important;
    }

    .stExpander {
        background: rgba(255,255,255,0.03) !important;
        border: 1px solid rgba(232,213,183,0.12) !important;
        border-radius: 10px !important;
    }

    .stProgress > div > div > div {
        background: linear-gradient(90deg, #c9a96e, #e8d5b7) !important;
    }

    label, .stMarkdown p {
        color: #b8aec4 !important;
    }

    .stSuccess {
        background: rgba(46,213,115,0.1) !important;
        border: 1px solid rgba(46,213,115,0.3) !important;
        border-radius: 8px !important;
    }

    .stWarning {
        background: rgba(255,165,0,0.1) !important;
        border: 1px solid rgba(255,165,0,0.3) !important;
        border-radius: 8px !important;
    }

    .stError {
        background: rgba(255,71,87,0.1) !important;
        border: 1px solid rgba(255,71,87,0.3) !important;
        border-radius: 8px !important;
    }

    .keyword-chip {
        display: inline-block;
        background: rgba(48,43,99,0.6);
        border: 1px solid rgba(201,169,110,0.25);
        color: #c9a96e;
        border-radius: 4px;
        padding: 0.15rem 0.55rem;
        font-size: 0.75rem;
        margin: 0.15rem;
        font-family: monospace;
    }

    .stat-box {
        background: rgba(255,255,255,0.04);
        border: 1px solid rgba(201,169,110,0.2);
        border-radius: 10px;
        padding: 1rem;
        text-align: center;
    }

    .stat-number {
        font-size: 1.8rem;
        font-weight: 700;
        color: #c9a96e;
        font-family: 'Playfair Display', serif;
    }

    .stat-label {
        font-size: 0.75rem;
        color: #7c6f8a;
        margin-top: 0.2rem;
        letter-spacing: 0.5px;
    }
</style>
""", unsafe_allow_html=True)


def search_arxiv(query, max_results=3):
    base_url = "http://export.arxiv.org/api/query"
    params = {
        "search_query": f"all:{query}",
        "start": 0,
        "max_results": max_results,
        "sortBy": "relevance",
        "sortOrder": "descending"
    }
    try:
        response = requests.get(base_url, params=params, timeout=15)
        response.raise_for_status()
        content = response.text
        entries = re.findall(r"<entry>(.*?)</entry>", content, re.DOTALL)
        results = []
        for entry in entries:
            title_match = re.search(r"<title>(.*?)</title>", entry, re.DOTALL)
            summary_match = re.search(r"<summary>(.*?)</summary>", entry, re.DOTALL)
            authors_match = re.findall(r"<name>(.*?)</name>", entry)
            published_match = re.search(r"<published>(.*?)</published>", entry)
            id_match = re.search(r"<id>(.*?)</id>", entry)
            if title_match and summary_match:
                arxiv_id = ""
                pdf_url = ""
                if id_match:
                    full_id = id_match.group(1).strip()
                    arxiv_id = full_id.split("/abs/")[-1]
                    pdf_url = f"https://arxiv.org/pdf/{arxiv_id}"
                published = ""
                if published_match:
                    published = published_match.group(1).strip()[:10]
                results.append({
                    "title": title_match.group(1).strip().replace("\n", " "),
                    "summary": summary_match.group(1).strip().replace("\n", " "),
                    "authors": ", ".join(authors_match[:3]) + (" et al." if len(authors_match) > 3 else ""),
                    "published": published,
                    "pdf_url": pdf_url,
                    "arxiv_id": arxiv_id
                })
        return results
    except Exception as e:
        st.error(f"arXiv API hatası: {str(e)}")
        return []


def translate_text(text, chunk_size=4000):
    if not text or not text.strip():
        return ""
    chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
    translated_parts = []
    for chunk in chunks:
        try:
            translated = GoogleTranslator(source="auto", target="tr").translate(chunk)
            translated_parts.append(translated if translated else chunk)
        except Exception:
            translated_parts.append(chunk)
    return " ".join(translated_parts)


def generate_summary(text, num_sentences=5):
    if not text or len(text.strip()) < 100:
        return text
    try:
        sentences = sent_tokenize(text)
        if len(sentences) <= num_sentences:
            return text
        try:
            stop_words = set(stopwords.words("english"))
        except Exception:
            stop_words = set()
        words = word_tokenize(text.lower())
        words = [w for w in words if w.isalpha() and w not in stop_words and len(w) > 3]
        word_freq = Counter(words)
        sentence_scores = {}
        for sent in sentences:
            sent_words = word_tokenize(sent.lower())
            score = sum(word_freq.get(w, 0) for w in sent_words if w.isalpha())
            if len(sent_words) > 0:
                sentence_scores[sent] = score / len(sent_words)
        top_sentences = sorted(sentence_scores, key=sentence_scores.get, reverse=True)[:num_sentences]
        ordered = [s for s in sentences if s in top_sentences]
        return " ".join(ordered)
    except Exception:
        words = text.split()
        return " ".join(words[:300]) + "..."


def extract_keywords(text, top_n=20):
    try:
        try:
            stop_words = set(stopwords.words("english"))
        except Exception:
            stop_words = set()
        tr_stops = {"ve", "bir", "bu", "da", "de", "ile", "için", "olan", "olan", "çok", "daha", "en"}
        stop_words.update(tr_stops)
        words = re.findall(r'\b[a-zA-ZğüşıöçĞÜŞİÖÇ]{4,}\b', text.lower())
        words = [w for w in words if w not in stop_words]
        freq = Counter(words)
        return freq.most_common(top_n)
    except Exception:
        return []


def detect_sections(text):
    section_patterns = {
        "Abstract": r"(?i)\babstract\b",
        "Introduction": r"(?i)\bintroduction\b",
        "Methodology": r"(?i)\b(methodology|methods|method)\b",
        "Results": r"(?i)\b(results|findings)\b",
        "Discussion": r"(?i)\bdiscussion\b",
        "Conclusion": r"(?i)\b(conclusion|conclusions)\b",
        "References": r"(?i)\breferences\b"
    }
    found_sections = {}
    sentences = text.split(". ")
    for section_name, pattern in section_patterns.items():
        for i, sentence in enumerate(sentences):
            if re.search(pattern, sentence):
                start = max(0, i)
                end = min(len(sentences), i + 30)
                found_sections[section_name] = ". ".join(sentences[start:end])
                break
    return found_sections


def estimate_reading_time(text):
    words = len(text.split())
    minutes = math.ceil(words / 200)
    return words, minutes


def generate_wordcloud_html(text):
    try:
        keywords = extract_keywords(text, top_n=40)
        if not keywords:
            return ""
        max_count = max(c for _, c in keywords) if keywords else 1
        colors = ["#c9a96e", "#e8d5b7", "#a07840", "#d4b896", "#8b6914",
                  "#f0e0c0", "#b8860b", "#daa520", "#cd853f", "#d2691e"]
        html = '<div style="background:#1a1230;border-radius:12px;padding:1.5rem;line-height:2.2;text-align:center;">'
        for i, (word, count) in enumerate(keywords):
            size = 0.75 + (count / max_count) * 1.5
            color = colors[i % len(colors)]
            opacity = 0.7 + (count / max_count) * 0.3
            html += f'<span style="font-size:{size:.2f}rem;color:{color};opacity:{opacity:.2f};margin:0.3rem 0.5rem;display:inline-block;font-weight:{"700" if size > 1.5 else "400"};">{word}</span>'
        html += '</div>'
        return html
    except Exception:
        return ""


def create_docx(title, original_text, translated_text, summary, source_type="API"):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title_para = doc.add_heading(level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.color.rgb = RGBColor(0x1a, 0x14, 0x25)
    run.font.size = Pt(18)

    doc.add_paragraph()

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(f"Kaynak: {source_type} | Akademik Makale Yönetim Sistemi")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x7c, 0x6f, 0x8a)

    doc.add_paragraph()
    doc.add_heading("📋 Rafine Özet (Türkçe)", level=1)
    summary_tr = translate_text(summary)
    sum_para = doc.add_paragraph(summary_tr if summary_tr else summary)
    sum_para.style.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_heading("🌐 Türkçe Çeviri", level=1)
    trans_para = doc.add_paragraph(translated_text)
    trans_para.style.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_heading("📄 Orijinal Metin (İngilizce)", level=1)
    orig_para = doc.add_paragraph(original_text[:5000] + ("..." if len(original_text) > 5000 else ""))
    orig_para.style.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def create_pdf_report(title, original_text, translated_text, summary, source_type="API"):
    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        pdf.set_font("Helvetica", "B", 16)
        safe_title = title.encode("latin-1", "replace").decode("latin-1")
        pdf.multi_cell(0, 10, safe_title, align="C")
        pdf.ln(3)

        pdf.set_font("Helvetica", "", 8)
        meta = f"Kaynak: {source_type} | Akademik Makale Yonetim Sistemi"
        pdf.cell(0, 6, meta, align="C")
        pdf.ln(6)
        pdf.set_draw_color(201, 169, 110)
        pdf.set_line_width(0.5)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(6)

        def add_section(heading, text, max_chars=4000):
            pdf.set_font("Helvetica", "B", 12)
            safe_heading = heading.encode("latin-1", "replace").decode("latin-1")
            pdf.cell(0, 8, safe_heading)
            pdf.ln(5)
            pdf.set_font("Helvetica", "", 10)
            safe_text = text[:max_chars].encode("latin-1", "replace").decode("latin-1")
            if len(text) > max_chars:
                safe_text += "..."
            pdf.multi_cell(0, 6, safe_text)
            pdf.ln(5)

        summary_tr = translate_text(summary)
        add_section("Rafine Ozet (Turkce)", summary_tr if summary_tr else summary)
        add_section("Turkce Ceviri", translated_text, max_chars=5000)
        add_section("Orijinal Metin (Ingilizce)", original_text, max_chars=3000)

        buf = io.BytesIO()
        pdf_bytes = pdf.output()
        buf.write(bytes(pdf_bytes))
        buf.seek(0)
        return buf
    except Exception as e:
        buf = io.BytesIO(b"PDF olusturulamadi.")
        buf.seek(0)
        return buf


def extract_pdf_text(uploaded_file):
    try:
        reader = PdfReader(uploaded_file)
        full_text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                full_text += page_text + "\n"
        return full_text.strip()
    except Exception as e:
        st.error(f"PDF okuma hatası: {str(e)}")
        return ""


st.markdown("""
<div class="main-header">
    <h1>📚 Akademik Makale Yönetim Sistemi</h1>
    <p>ARAŞTIR · ÇEVİR · ÖZETLE · İNDİR</p>
</div>
<div class="divider-line"></div>
""", unsafe_allow_html=True)

tab1, tab2 = st.tabs(["🔍 API ile Makale Ara", "📂 PDF Yükle ve İşle"])


with tab1:
    st.markdown('<div style="height:0.5rem"></div>', unsafe_allow_html=True)

    col_search, col_count = st.columns([3, 1])
    with col_search:
        query = st.text_input(
            "Anahtar Kelime",
            placeholder="örn: machine learning, transformer, neural network...",
            label_visibility="collapsed"
        )
    with col_count:
        max_results = st.slider("Makale Sayısı", min_value=1, max_value=5, value=3, label_visibility="collapsed")

    search_col, _ = st.columns([1, 4])
    with search_col:
        search_btn = st.button("🔍 Ara", use_container_width=True)

    if search_btn and query:
        with st.spinner("arXiv veritabanında aranıyor..."):
            papers = search_arxiv(query, max_results)

        if not papers:
            st.warning("Sonuç bulunamadı. Farklı anahtar kelimeler deneyin.")
        else:
            st.markdown(f'<div class="section-label">{len(papers)} makale bulundu</div>', unsafe_allow_html=True)

            for idx, paper in enumerate(papers):
                with st.spinner(f"[{idx+1}/{len(papers)}] Çeviriliyor: {paper['title'][:50]}..."):
                    translated_summary = translate_text(paper["summary"])
                    refined_summary = generate_summary(paper["summary"], num_sentences=4)
                    refined_tr = translate_text(refined_summary)

                st.markdown(f"""
                <div class="card">
                    <div class="card-title">{paper['title']}</div>
                    <div class="card-meta">
                        👤 {paper['authors']} &nbsp;|&nbsp; 📅 {paper['published']}
                    </div>
                </div>
                """, unsafe_allow_html=True)

                c1, c2, c3 = st.columns(3)

                with c1:
                    st.markdown('<div class="section-label">📄 Orijinal Özet</div>', unsafe_allow_html=True)
                    st.markdown(f'<div class="text-block">{paper["summary"][:600]}{"..." if len(paper["summary"]) > 600 else ""}</div>', unsafe_allow_html=True)

                with c2:
                    st.markdown('<div class="section-label">🌐 Türkçe Çeviri</div>', unsafe_allow_html=True)
                    st.markdown(f'<div class="text-block">{translated_summary[:600]}{"..." if len(translated_summary) > 600 else ""}</div>', unsafe_allow_html=True)

                with c3:
                    st.markdown('<div class="section-label">✨ Rafine Özet (TR)</div>', unsafe_allow_html=True)
                    st.markdown(f'<div class="text-block">{refined_tr[:600]}{"..." if len(refined_tr) > 600 else ""}</div>', unsafe_allow_html=True)

                st.markdown('<div style="height:0.6rem"></div>', unsafe_allow_html=True)
                dl1, dl2, dl3 = st.columns([1, 1, 4])

                safe_title = re.sub(r'[^\w\s-]', '', paper['title'])[:40].strip().replace(" ", "_")

                with dl1:
                    pdf_buf = create_pdf_report(
                        paper["title"],
                        paper["summary"],
                        translated_summary,
                        refined_summary,
                        source_type="arXiv"
                    )
                    st.download_button(
                        label="📥 PDF İndir",
                        data=pdf_buf,
                        file_name=f"{safe_title}.pdf",
                        mime="application/pdf",
                        key=f"pdf_{idx}"
                    )

                with dl2:
                    docx_buf = create_docx(
                        paper["title"],
                        paper["summary"],
                        translated_summary,
                        refined_summary,
                        source_type="arXiv"
                    )
                    st.download_button(
                        label="📝 DOCX İndir",
                        data=docx_buf,
                        file_name=f"{safe_title}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"docx_{idx}"
                    )

                st.markdown('<hr style="border:none;border-top:1px solid rgba(232,213,183,0.08);margin:1.2rem 0;">', unsafe_allow_html=True)

    elif search_btn and not query:
        st.warning("Lütfen bir anahtar kelime girin.")


with tab2:
    st.markdown('<div style="height:0.5rem"></div>', unsafe_allow_html=True)

    uploaded_file = st.file_uploader(
        "PDF dosyanızı buraya sürükleyin veya seçin",
        type=["pdf"],
        label_visibility="collapsed"
    )

    if uploaded_file is not None:
        st.markdown(f'<div class="info-pill">📄 {uploaded_file.name}</div><div class="info-pill">💾 {round(uploaded_file.size/1024, 1)} KB</div>', unsafe_allow_html=True)

        with st.spinner("PDF okunuyor..."):
            raw_text = extract_pdf_text(uploaded_file)

        if not raw_text:
            st.error("PDF'den metin ayıklanamadı. Dosyanın metin katmanı içerdiğinden emin olun.")
        else:
            word_count, read_minutes = estimate_reading_time(raw_text)
            char_count = len(raw_text)
            sections = detect_sections(raw_text)

            m1, m2, m3, m4 = st.columns(4)
            with m1:
                st.markdown(f'<div class="stat-box"><div class="stat-number">{word_count:,}</div><div class="stat-label">KELİME</div></div>', unsafe_allow_html=True)
            with m2:
                st.markdown(f'<div class="stat-box"><div class="stat-number">{char_count:,}</div><div class="stat-label">KARAKTER</div></div>', unsafe_allow_html=True)
            with m3:
                st.markdown(f'<div class="stat-box"><div class="stat-number">{read_minutes}</div><div class="stat-label">DAKİKA OKUMA</div></div>', unsafe_allow_html=True)
            with m4:
                st.markdown(f'<div class="stat-box"><div class="stat-number">{len(sections)}</div><div class="stat-label">BÖLÜM TESPİT</div></div>', unsafe_allow_html=True)

            st.markdown('<div style="height:1rem"></div>', unsafe_allow_html=True)

            if sections:
                st.markdown('<div class="section-label">📑 Tespit Edilen Bölümler</div>', unsafe_allow_html=True)
                section_html = "".join([f'<span class="badge">✓ {s}</span>' for s in sections.keys()])
                st.markdown(section_html, unsafe_allow_html=True)

            st.markdown('<div style="height:0.8rem"></div>', unsafe_allow_html=True)

            keywords = extract_keywords(raw_text, top_n=25)
            if keywords:
                st.markdown('<div class="section-label">🔑 Öne Çıkan Anahtar Kelimeler</div>', unsafe_allow_html=True)
                kw_html = "".join([f'<span class="keyword-chip">{word} <span style="opacity:0.5">({count})</span></span>' for word, count in keywords])
                st.markdown(kw_html, unsafe_allow_html=True)

            st.markdown('<div style="height:0.8rem"></div>', unsafe_allow_html=True)

            with st.expander("☁️ Kelime Bulutu Göster", expanded=False):
                wc_html = generate_wordcloud_html(raw_text)
                if wc_html:
                    st.markdown(wc_html, unsafe_allow_html=True)

            st.markdown('<div style="height:0.8rem"></div>', unsafe_allow_html=True)

            process_btn = st.button("⚙️ Çevir ve Özetle", use_container_width=False)

            if process_btn:
                progress_bar = st.progress(0)
                status_text = st.empty()

                status_text.markdown('<div class="section-label">Metin çevriliyor...</div>', unsafe_allow_html=True)
                progress_bar.progress(20)
                translated_text = translate_text(raw_text)
                progress_bar.progress(60)

                status_text.markdown('<div class="section-label">Özet çıkarılıyor...</div>', unsafe_allow_html=True)
                summary_en = generate_summary(raw_text, num_sentences=6)
                progress_bar.progress(75)

                summary_tr = translate_text(summary_en)
                progress_bar.progress(90)

                status_text.empty()
                progress_bar.progress(100)
                progress_bar.empty()

                st.success("İşlem tamamlandı!")
                st.markdown('<div style="height:0.5rem"></div>', unsafe_allow_html=True)

                r1, r2 = st.columns(2)

                with r1:
                    st.markdown('<div class="section-label">✨ Türkçe Özet</div>', unsafe_allow_html=True)
                    st.markdown(f'<div class="text-block">{summary_tr}</div>', unsafe_allow_html=True)

                with r2:
                    st.markdown('<div class="section-label">🌐 Türkçe Çeviri (ilk 1000 karakter)</div>', unsafe_allow_html=True)
                    st.markdown(f'<div class="text-block">{translated_text[:1000]}{"..." if len(translated_text) > 1000 else ""}</div>', unsafe_allow_html=True)

                if sections:
                    st.markdown('<div style="height:0.8rem"></div>', unsafe_allow_html=True)
                    st.markdown('<div class="section-label">📑 Bölüm Bazlı Çeviriler</div>', unsafe_allow_html=True)
                    for sec_name, sec_text in list(sections.items())[:4]:
                        with st.expander(f"📖 {sec_name}"):
                            sec_tr = translate_text(sec_text[:2000])
                            st.markdown(f'<div class="text-block">{sec_tr}</div>', unsafe_allow_html=True)

                st.markdown('<div style="height:0.8rem"></div>', unsafe_allow_html=True)
                doc_title = uploaded_file.name.replace(".pdf", "")
                safe_name = re.sub(r'[^\w\s-]', '', doc_title)[:40].strip().replace(" ", "_")

                dl_c1, dl_c2 = st.columns([1, 1])
                with dl_c1:
                    pdf_out = create_pdf_report(
                        doc_title,
                        raw_text,
                        translated_text,
                        summary_en,
                        source_type="Yerel PDF"
                    )
                    st.download_button(
                        label="📥 PDF Olarak İndir",
                        data=pdf_out,
                        file_name=f"{safe_name}_islenmis.pdf",
                        mime="application/pdf",
                        key="local_pdf"
                    )

                with dl_c2:
                    docx_out = create_docx(
                        doc_title,
                        raw_text,
                        translated_text,
                        summary_en,
                        source_type="Yerel PDF"
                    )
                    st.download_button(
                        label="📝 DOCX Olarak İndir",
                        data=docx_out,
                        file_name=f"{safe_name}_islenmis.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="local_docx"
                    )

    else:
        st.markdown("""
        <div style="text-align:center;padding:3rem 0;color:#7c6f8a;">
            <div style="font-size:3rem;margin-bottom:1rem;">📂</div>
            <div style="font-size:1rem;font-weight:500;color:#9b8ea0;">PDF dosyanızı yükleyin</div>
            <div style="font-size:0.85rem;margin-top:0.5rem;">Metin katmanı içeren akademik makaleler desteklenmektedir</div>
        </div>
        """, unsafe_allow_html=True)

st.markdown("""
<div style="text-align:center;padding:2rem 0 1rem 0;color:#4a3f5c;font-size:0.75rem;letter-spacing:1px;">
    AKADEMIK MAKALE YÖNETİM SİSTEMİ &nbsp;·&nbsp; arXiv API &nbsp;·&nbsp; Google Translate &nbsp;·&nbsp; NLTK
</div>
""", unsafe_allow_html=True)
